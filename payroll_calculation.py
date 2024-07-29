import pandas as pd
import time
import os
from docx import Document
from fpdf import FPDF

# Load the dataset
df = pd.read_excel("payroll_dataset.xlsx")

# Function to calculate payroll for each employee and convert to rupees per month
def calculate_payroll(row, exchange_rate):
    # Perform payroll calculation based on attributes in the dataset
    # For demonstration, let's calculate total pay by adding pay rate, bonus, and commission
    total_pay_dollars = row["Pay Rate"] + row["Bonus"] + row["Commission"]
    total_pay_rupees_per_month = total_pay_dollars * exchange_rate
    return total_pay_rupees_per_month

# Assuming exchange rate from USD to INR
exchange_rate = 74.25  # Change this to your actual exchange rate

# Calculate payroll for each employee in rupees per month
df["Total Pay (INR)"] = df.apply(calculate_payroll, axis=1, args=(exchange_rate,))

# Generate Word document content with employee name and total pay in rupees per month
doc = Document()
doc.add_heading('Payroll Summary', level=1)
for index, row in df.iterrows():
    doc.add_paragraph(f"{row['Full Name']}: ₹{row['Total Pay (INR)']:.2f}")

# Save the Word document
doc_file_path = "payroll_summary.docx"
doc.save(doc_file_path)

print(f"Payroll summary saved as '{doc_file_path}'")

# Convert Word document to PDF
def convert_to_pdf(input_path, output_path):
    doc = Document(input_path)
    pdf = FPDF()
    pdf.add_page()
    for para in doc.paragraphs:
        text = para.text.encode("latin-1", errors='replace').decode('latin-1')
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=text)
    pdf.output(output_path)

# Convert Word document to PDF
pdf_file_path = "payroll_summary.pdf"
convert_to_pdf(doc_file_path, pdf_file_path)

print(f"Payroll summary saved as '{pdf_file_path}'")

# Close Microsoft Word if it is already open
os.system("TASKKILL /F /IM WINWORD.EXE")

# Open the Word document
os.system(f"start WINWORD.EXE {doc_file_path}")

# Sleep to allow Word to open
time.sleep(5)

# Close Word after a few seconds
os.system("TASKKILL /F /IM WINWORD.EXE")

# Move the PDF file to the current directory
current_directory = os.path.dirname(os.path.abspath(__file__))
os.rename(pdf_file_path, os.path.join(current_directory, pdf_file_path))
